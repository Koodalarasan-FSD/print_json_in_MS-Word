from flask import Flask, send_file
from docx import Document
import os

# Type(Paste) http://127.0.0.1:5000/generate_word in browser

app = Flask(__name__)

@app.route('/generate_word')
def generate_word():
    # Specify the directory where you want to save the file
    save_directory = '/path/to/your/directory'

    # Create the specified directory if it doesn't exist
    if not os.path.exists(save_directory):
        os.makedirs(save_directory)

    # Create a new Word document
    doc = Document()

    # Sample JSON data
    json_data = {
        "name": "John Doe",
        "age": 30,
        "city": "Example City"
    }

    # Add content to the Word document
    doc.add_heading('JSON Data', level=1)
    for key, value in json_data.items():
        doc.add_paragraph(f"{key}: {value}")

    # Save the Word document in the specified directory
    save_path = os.path.join(save_directory, 'output.docx')
    doc.save(save_path)

    #print(save_path)
    #print(save_directory)
    
    # Send the Word document as a response
    return send_file(save_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
